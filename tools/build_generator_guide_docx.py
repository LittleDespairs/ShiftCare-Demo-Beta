from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCES = [
    ROOT / "docs" / "GENERATOR_GUIDE_USERS_RU.md",
    ROOT / "docs" / "GENERATOR_GUIDE_TECH_RU.md",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, rows):
    table = document.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F6F8FA")
    document.add_paragraph()


def add_markdown_paragraph(document, line):
    stripped = line.strip()
    if not stripped:
        return

    if stripped.startswith("**") and stripped.endswith("**"):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(stripped.strip("*"))
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        return

    paragraph = document.add_paragraph()
    if stripped[0:3].replace(".", "").isdigit() and ". " in stripped[:4]:
        paragraph.style = "List Number"
        stripped = stripped.split(". ", 1)[1]
    elif stripped.startswith("- "):
        paragraph.style = "List Bullet"
        stripped = stripped[2:]

    run = paragraph.add_run(stripped)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def build_docx(source: Path):
    output = source.with_suffix(".docx")
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size, color in [
        ("Title", 22, RGBColor(31, 78, 121)),
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(47, 84, 150)),
        ("Heading 3", 11, RGBColor(31, 78, 121)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(line[2:].strip())
            index += 1
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
            index += 1
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            index += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                candidate = lines[index].strip()
                if not set(candidate.replace("|", "").replace("-", "").replace(":", "").strip()) <= set():
                    table_lines.append(candidate)
                index += 1
            rows = []
            for table_line in table_lines:
                values = [cell.strip() for cell in table_line.strip("|").split("|")]
                if all(value.replace("-", "").replace(":", "").strip() == "" for value in values):
                    continue
                rows.append(values)
            if rows:
                add_table(document, rows)
            continue

        add_markdown_paragraph(document, line)
        index += 1

    document.core_properties.title = "Руководство по работе генератора расписания"
    document.core_properties.subject = "ShiftCare"
    document.core_properties.author = "Codex"
    document.save(output)
    return output


if __name__ == "__main__":
    for source_path in SOURCES:
        print(build_docx(source_path))
